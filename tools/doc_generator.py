from docx import Document
from docx.shared import Pt
import os

def create_admission_report(university, program, content, output_path):
    doc = Document()
    
    # Styled Title
    title = doc.add_heading(f'Admissions Strategy: {university}', 0)
    doc.add_heading(f'Program: {program}', level=1)
    
    # Split content by line to handle Markdown-ish formatting
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Basic Markdown Handling
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('*') or line.startswith('-'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    doc.save(output_path)
    return output_path